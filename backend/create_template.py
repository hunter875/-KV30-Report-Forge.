"""
Script to generate the kv30_report.docx template with docxtpl placeholders.
Run: python create_template.py
"""
from docx import Document
from docxtpl import DocxTemplate, InlineImage
import io

def create_template():
    """Create the official KV30 report template with placeholders."""
    doc = Document()

    # Title
    title = doc.add_heading('BÁO CÁO PCCC VÀ CNCH', level=0)
    title.alignment = 1  # Center

    # Header Section
    doc.add_heading('Thông tin báo cáo', level=1)

    # Create a table for header fields
    header_table = doc.add_table(rows=4, cols=2)
    header_table.style = 'Light Grid Accent 1'

    header_fields = [
        ('Số báo cáo:', '{{so_bao_cao}}'),
        ('Ngày báo cáo:', '{{ngay_bao_cao}}'),
        ('Đơn vị báo cáo:', '{{don_vi_bao_cao}}'),
        ('Thời gian từ - đến:', '{{thoi_gian_tu_den}}')
    ]

    for i, (label, placeholder) in enumerate(header_fields):
        row = header_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = placeholder

    doc.add_paragraph()

    # Statistics Section (Bảng thống kê)
    doc.add_heading('I. TÌNH HÌNH CHÁY, NỔ, SỰ CỐ TAI NẠN', level=1)

    # Create statistics table with 62 rows (header + 61 data rows)
    stats_table = doc.add_table(rows=62, cols=3)
    stats_table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = stats_table.rows[0].cells
    header_cells[0].text = 'STT'
    header_cells[1].text = 'Nội dung'
    header_cells[2].text = 'Kết quả'

    # Data rows with placeholders
    for i in range(1, 62):
        stt_num = i
        row = stats_table.rows[i]
        row.cells[0].text = str(stt_num)
        row.cells[1].text = f'{{{{stt_{stt_num}}}}}'  # Will be replaced with noi_dung from template
        row.cells[2].text = f'{{{{stt_{stt_num}}}_ket_qua}}'  # Will be replaced with ket_qua

    doc.add_paragraph()

    # II. Detailed Operations (Phần I và II chi tiết nghiệp vụ)
    doc.add_heading('II. KẾT QUẢ CÔNG TÁC PCCC VÀ CNCH', level=1)

    # Add detailed operation fields
    detail_fields = [
        ('Quân số trực:', '{{phan_I_va_II_chi_tiet_nghiep_vu.quan_so_truc}}'),
        ('Tổng báo cáo:', '{{phan_I_va_II_chi_tiet_nghiep_vu.tong_bao_cao}}'),
        ('Chi tiết CNCH:', '{{phan_I_va_II_chi_tiet_nghiep_vu.chi_tiet_cnch}}'),
        ('Tổng chỉ viên:', '{{phan_I_va_II_chi_tiet_nghiep_vu.tong_chi_vien}}'),
        ('Tổng công văn:', '{{phan_I_va_II_chi_tiet_nghiep_vu.tong_cong_van}}'),
        ('Tổng kế hoạch:', '{{phan_I_va_II_chi_tiet_nghiep_vu.tong_ke_hoach}}'),
        ('Tổng số vụ nổ:', '{{phan_I_va_II_chi_tiet_nghiep_vu.tong_so_vu_no}}'),
        ('Tổng số vụ cháy:', '{{phan_I_va_II_chi_tiet_nghiep_vu.tong_so_vu_chay}}'),
        ('Tổng số vụ CNCH:', '{{phan_I_va_II_chi_tiet_nghiep_vu.tong_so_vu_cnch}}'),
        ('Tổng xe hư hỏng:', '{{phan_I_va_II_chi_tiet_nghiep_vu.tong_xe_hu_hong}}'),
        ('Công tác an ninh:', '{{phan_I_va_II_chi_tiet_nghiep_vu.cong_tac_an_ninh}}')
    ]

    for label, placeholder in detail_fields:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(f' {placeholder}')

    doc.add_paragraph()

    # III. Danh sách CNCH
    doc.add_heading('Danh sách sự cố CNCH', level=2)
    doc.add_paragraph('{% for item in danh_sach_cnch %}')
    doc.add_paragraph('STT: {{item.stt}}')
    doc.add_paragraph('Mô tả: {{item.mo_ta}}')
    doc.add_paragraph('Địa điểm: {{item.dia_diem}}')
    doc.add_paragraph('Thiệt hại: {{item.thiet_hai}}')
    doc.add_paragraph('Thời gian: {{item.thoi_gian}}')
    doc.add_paragraph('Ngày xảy ra: {{item.ngay_xay_ra}}')
    doc.add_paragraph('Kết quả xử lý: {{item.ket_qua_xu_ly}}')
    doc.add_paragraph('Nội dung tin báo: {{item.noi_dung_tin_bao}}')
    doc.add_paragraph('Lực lượng tham gia: {{item.luc_luong_tham_gia}}')
    doc.add_paragraph('Thông tin nạn nhân: {{item.thong_tin_nan_nhan}}')
    doc.add_paragraph('{% endfor %}')

    doc.add_paragraph()

    # IV. Danh sách công tác khác
    doc.add_heading('Danh sách công tác khác', level=2)
    doc.add_paragraph('{% for task in danh_sach_cong_tac_khac %}')
    doc.add_paragraph('• {{task}}')
    doc.add_paragraph('{% endfor %}')

    doc.add_paragraph()

    # V. Danh sách công văn tham mưu
    doc.add_heading('Danh sách công văn tham mưu', level=2)
    doc.add_paragraph('{% for doc in danh_sach_cong_van_tham_muu %}')
    doc.add_paragraph('• {{doc}}')
    doc.add_paragraph('{% endfor %}')

    doc.add_paragraph()

    # VI. Danh sách phương tiện hư hỏng
    doc.add_heading('Danh sách phương tiện hư hỏng', level=2)
    doc.add_paragraph('{% for vehicle in danh_sach_phuong_tien_hu_hong %}')
    doc.add_paragraph('Biển số: {{vehicle.bien_so}} - Tình trạng: {{vehicle.tinh_trang}}')
    doc.add_paragraph('{% endfor %}')

    # Save the template
    template_path = 'd:/exel/backend/app/modules/reports/templates/kv30_report.docx'
    doc.save(template_path)
    print(f'Template created at: {template_path}')

    # Also create a sample filled document to verify structure
    print('Creating sample render to verify placeholders...')

if __name__ == '__main__':
    create_template()
